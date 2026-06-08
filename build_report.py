# -*- coding: utf-8 -*-
"""
Сборка итогового отчёта по курсовому проекту «Базы данных» в .docx.
Запуск:  python build_report.py
Результат: Отчет_БД_anywai.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
SQL_DIR = ROOT / "sql"
OUT = ROOT / "Отчет_БД_anywai.docx"
ER_IMAGE = ROOT / "er_diagram.png"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
BODY_SIZE = 14
CODE_SIZE = 9


# ─────────────────────────── низкоуровневые помощники ───────────────────────
def set_cell_background(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)


def add_body(doc: Document, text: str, *, justify: bool = True, italic: bool = False,
             bold: bool = False, first_line_indent: bool = True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def add_h1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs:
        r.font.name = BODY_FONT
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_h2(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = BODY_FONT
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_code(doc: Document, code: str, caption: str | None = None):
    """Моноширинный блок кода в одну ячейку таблицы с серой заливкой."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F2F2F2")
    # бордюр
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "BBBBBB")
        borders.append(e)
    tbl_pr.append(borders)

    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in code.rstrip("\n").split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else "\u00A0")
        run.font.name = MONO_FONT
        run.font.size = Pt(CODE_SIZE)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), MONO_FONT)
        rfonts.set(qn("w:hAnsi"), MONO_FONT)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(12)
    return table


def add_bullets(doc: Document, items: list[str]):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc: Document, items: list[str]):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_table_grid(doc: Document, header: list[str], rows: list[list[str]],
                   widths: list[float] | None = None, caption: str | None = None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, text in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = BODY_FONT
        set_cell_background(hdr[i], "D9E2F3")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(12)
            run.font.name = BODY_FONT
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(12)
    return table


def read_sql(name: str) -> str:
    return (SQL_DIR / name).read_text(encoding="utf-8")


# ─────────────────────────────── титульный лист ─────────────────────────────
def title_page(doc: Document) -> None:
    def c(text, *, size=14, bold=False, italic=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = BODY_FONT
        return p

    c("МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", size=12)
    c("ФГБОУ ВО «РГПУ им. А.И. ГЕРЦЕНА»", size=12, space_after=24)
    c("Направление подготовки", size=12, space_before=12)
    c("09.03.01 – Информатика и вычислительная техника", size=12)
    c("Профиль «Технологии разработки программного обеспечения "
      "и обработки больших данных»", size=12, space_after=60)

    c("Отчет по проекту", size=16, bold=True, space_before=24)
    c("«Проектирование базы данных истории диалогов и аналитики "
      "AI-чат-сервиса»", size=16, bold=True, space_after=72)

    c("Работу выполнил:", size=14, space_before=24)
    c("Студент 2 курса, ИВТ-1.1", size=14)
    c("Шарманов Даниил Андреевич", size=14, space_after=18)
    c("Преподаватель:", size=14)
    c("Жуков Николай Николаевич", size=14, space_after=120)

    c("Санкт-Петербург", size=12, space_before=48)
    c("2026", size=12)
    doc.add_page_break()


# ──────────────────────────────── разделы ───────────────────────────────────
def section_intro(doc: Document) -> None:
    add_h1(doc, "Введение")
    add_body(doc,
        "Современные сервисы на основе больших языковых моделей (LLM) накапливают "
        "значительные объёмы данных: историю переписки пользователей, метаданные "
        "запросов к моделям, события поведения. Грамотно спроектированная база данных "
        "позволяет хранить эти данные целостно, быстро выполнять аналитические запросы "
        "и масштабировать сервис без потери качества.")
    add_body(doc,
        "Объект проекта — реальный AI-сервис anywai (Telegram-бот с веб-клиентом), "
        "взаимодействующий с моделями семейств Gemini, Claude, GPT и Llama. Полная "
        "схема сервиса включает десятки таблиц (биллинг, подписки, реферальная "
        "программа, бизнес-режим). Для курсового проекта выделен самостоятельный, "
        "логически завершённый поддомен — «История диалогов и аналитика поведения "
        "пользователей». Это позволяет сосредоточиться на проектировании БД, не "
        "распыляясь на прикладную логику сервиса.")
    add_body(doc,
        "Цель проекта — спроектировать нормализованную реляционную базу данных для "
        "выбранной предметной области, реализовать её средствами СУБД PostgreSQL, "
        "снабдить индексами, триггерами, функциями и хранимыми процедурами, а также "
        "оценить целесообразность применения NoSQL-решений.")
    add_body(doc, "Для достижения цели поставлены задачи:", first_line_indent=False)
    add_numbered(doc, [
        "Описать предметную область и определить состав сущностей.",
        "Обосновать выбор системы управления базами данных.",
        "Провести нормализацию схемы до третьей нормальной формы (3НФ).",
        "Построить ER-диаграмму в нотации IDEF1x.",
        "Разработать физическую модель (DDL) с ограничениями целостности.",
        "Создать индексы под типовые поисковые запросы и обосновать их.",
        "Реализовать триггеры, функции и хранимые процедуры.",
        "Оценить целесообразность перехода на NoSQL.",
        "Реализовать CRUD-взаимодействие с БД через REST API.",
    ])


def section_subject(doc: Document) -> None:
    add_h1(doc, "1. Предметная область")
    add_h2(doc, "1.1. Описание предметной области")
    add_body(doc,
        "AI-чат-сервис предоставляет пользователям доступ к нескольким языковым "
        "моделям через единый интерфейс. Пользователь авторизуется через Telegram, "
        "ведёт один или несколько диалогов, отправляет текстовые сообщения и "
        "вложения (изображения, документы), а сервис возвращает ответы выбранной "
        "модели. Каждое значимое действие фиксируется в журнале событий для "
        "последующей продуктовой аналитики.")
    add_body(doc, "В базе данных хранятся:", first_line_indent=False)
    add_bullets(doc, [
        "учётные данные пользователей (идентификатор Telegram, имя, язык);",
        "справочники провайдеров и моделей LLM;",
        "диалоги пользователей и сообщения внутри них;",
        "метаданные вложений (без самих файлов);",
        "свёртки (краткие пересказы) длинных диалогов;",
        "журнал аналитических событий с произвольными свойствами.",
    ])
    add_body(doc,
        "Сознательно НЕ моделируются (вынесены за границы проекта): биллинг и "
        "оплаты, премиум-подписки и ключи активации, реферальная программа, "
        "бизнес-режим Telegram. Эти поддомены образуют отдельную предметную "
        "область и не нужны для задачи хранения истории и аналитики.")

    add_h2(doc, "1.2. Перечень сущностей")
    add_table_grid(doc,
        ["Сущность", "Назначение"],
        [
            ["users", "Пользователи сервиса; один пользователь соответствует одному Telegram-аккаунту."],
            ["llm_providers", "Справочник провайдеров моделей (Google, Anthropic, OpenAI, Groq)."],
            ["llm_models", "Справочник конкретных LLM-моделей."],
            ["dialogues", "Диалоги (чаты) пользователя."],
            ["messages", "Сообщения внутри диалога с привязкой к роли и модели."],
            ["attachments", "Метаданные файлов, прикреплённых к сообщению."],
            ["dialogue_summaries", "Свёртки старых сообщений длинных диалогов (1:1 к диалогу)."],
            ["event_kinds", "Справочник типов аналитических событий."],
            ["analytics_events", "Append-only журнал поведенческих и системных событий."],
        ],
        widths=[4.5, 11.5],
        caption="Таблица 1. Сущности предметной области")

    add_h2(doc, "1.3. Бизнес-правила")
    add_numbered(doc, [
        "Каждый диалог принадлежит ровно одному пользователю; при удалении "
        "пользователя его диалоги удаляются каскадно.",
        "Сообщение имеет роль: user (пользователь), assistant (модель) или "
        "system (системное).",
        "Только сообщение роли assistant связано с моделью; у user/system модели нет.",
        "Один диалог может иметь не более одной активной свёртки.",
        "Журнал событий неизменяем: записи нельзя удалять и редактировать "
        "(гарантируется триггером; допускается лишь системное обнуление ссылок "
        "при удалении связанных сущностей).",
        "Метаданные вложения всегда привязаны к конкретному сообщению.",
    ])


def section_dbms(doc: Document) -> None:
    add_h1(doc, "2. Обоснование выбора СУБД")
    add_body(doc,
        "В качестве СУБД выбрана PostgreSQL 16 — свободная объектно-реляционная "
        "система. Выбор обусловлен характером предметной области: данные сильно "
        "связаны (пользователи → диалоги → сообщения → вложения), требуют "
        "целостности на уровне внешних ключей и активно используются в "
        "аналитических запросах с агрегацией.")
    add_body(doc, "Ключевые причины выбора PostgreSQL:", first_line_indent=False)
    add_bullets(doc, [
        "Полная поддержка ACID-транзакций и внешних ключей — критично для "
        "целостности связанных данных.",
        "Тип JSONB с индексацией GIN — позволяет гибко хранить разнородные "
        "свойства событий, не плодя таблицы под каждый тип.",
        "Богатый набор индексов: B-tree, частичные (partial), выражений, GIN — "
        "под разные сценарии поиска.",
        "Оконные функции, CTE, GROUPING SETS/ROLLUP, LATERAL — мощный "
        "аналитический SQL «из коробки».",
        "Серверные функции и процедуры на PL/pgSQL, триггеры — бизнес-логика "
        "и инварианты на уровне БД.",
        "Зрелость, надёжность, бесплатность и совместимость с исходным "
        "сервисом anywai (он уже работает на PostgreSQL).",
    ])
    add_body(doc, "Сравнение с альтернативами приведено в таблице 2.",
             first_line_indent=False)
    add_table_grid(doc,
        ["СУБД", "Почему не выбрана для данной задачи"],
        [
            ["MySQL/MariaDB", "Слабее поддержка JSON и оконных функций в старых версиях; "
                              "беднее экосистема расширений и типов."],
            ["SQLite", "Файловая встраиваемая БД; не рассчитана на параллельную запись "
                       "многих пользователей и серверную нагрузку."],
            ["MongoDB", "Документная модель; нет внешних ключей и полноценных JOIN, "
                        "что усложняет аналитику по связанным сущностям."],
            ["Oracle/MS SQL", "Коммерческие лицензии; избыточны для учебного и "
                              "среднемасштабного проекта."],
        ],
        widths=[4.0, 12.0],
        caption="Таблица 2. Сравнение СУБД")


def section_norm(doc: Document) -> None:
    add_h1(doc, "3. Нормализация базы данных")
    add_body(doc,
        "Нормализация — процесс приведения схемы к виду, исключающему избыточность "
        "и аномалии вставки, обновления и удаления. Ниже показан путь от "
        "ненормализованного представления к третьей нормальной форме на примере "
        "данных о сообщениях.")

    add_h2(doc, "3.1. Исходное (ненормализованное) представление")
    add_body(doc,
        "Допустим, всю информацию о переписке хранили в одной «широкой» таблице:",
        first_line_indent=False)
    add_code(doc,
        "chat_log(\n"
        "    message_id, telegram_id, username, first_name, language,\n"
        "    dialogue_title, message_text, role,\n"
        "    model_code, model_name, provider_name, is_premium,\n"
        "    attachments_list, created_at\n"
        ")",
        caption="Листинг 1. Ненормализованная таблица")
    add_body(doc,
        "Такая таблица содержит повторяющиеся группы (список вложений в одном "
        "поле), а также многократно дублирует данные пользователя и модели в "
        "каждой строке. Это порождает аномалии: при переименовании модели "
        "нужно править тысячи строк (аномалия обновления); нельзя завести "
        "модель, которую ещё не использовали (аномалия вставки); удаление "
        "последнего сообщения с моделью теряет сведения о самой модели "
        "(аномалия удаления).")

    add_h2(doc, "3.2. Первая нормальная форма (1НФ)")
    add_body(doc,
        "Требование 1НФ: все атрибуты атомарны, нет повторяющихся групп. "
        "Поле attachments_list со списком файлов нарушает 1НФ. Выносим вложения "
        "в отдельную таблицу attachments со связью к сообщению, а каждое значение "
        "делаем атомарным.")
    add_body(doc,
        "Примечание о JSONB: поле properties_json в журнале событий формально "
        "напоминает нарушение 1НФ, однако это осознанное проектное решение для "
        "слабоструктурированных данных (см. раздел 7.4), а не повторяющаяся "
        "группа реляционных атрибутов.")

    add_h2(doc, "3.3. Вторая нормальная форма (2НФ)")
    add_body(doc,
        "Требование 2НФ: таблица в 1НФ, и каждый неключевой атрибут зависит от "
        "всего первичного ключа, а не от его части. Актуально при составных "
        "ключах. Разделяя данные на сущности с суррогатными первичными ключами "
        "(id), мы устраняем частичные зависимости: атрибуты пользователя "
        "зависят только от users.id, атрибуты модели — только от llm_models.id, "
        "атрибуты сообщения — только от messages.id.")

    add_h2(doc, "3.4. Третья нормальная форма (3НФ)")
    add_body(doc,
        "Требование 3НФ: таблица во 2НФ, и нет транзитивных зависимостей "
        "неключевых атрибутов. В исходной таблице provider_name транзитивно "
        "зависел от model_code (model_code → provider_name), а не напрямую от "
        "ключа сообщения. Выносим провайдеров в llm_providers, модели — в "
        "llm_models со ссылкой provider_id. Аналогично имя пользователя и язык "
        "выносятся в users.")
    add_body(doc,
        "Итоговая схема из девяти таблиц находится в 3НФ: каждый неключевой "
        "атрибут зависит только от первичного ключа своей таблицы. Соответствие "
        "исходных полей нормализованным таблицам приведено в таблице 3.")
    add_table_grid(doc,
        ["Поле исходной таблицы", "Куда вынесено", "Причина"],
        [
            ["username, first_name, language", "users", "зависят от пользователя, не от сообщения"],
            ["dialogue_title", "dialogues", "зависит от диалога"],
            ["model_name, is_premium", "llm_models", "зависят от модели"],
            ["provider_name", "llm_providers", "транзитивная зависимость через модель"],
            ["attachments_list", "attachments", "повторяющаяся группа (1НФ)"],
            ["message_text, role, created_at", "messages", "собственные атрибуты сообщения"],
        ],
        widths=[5.0, 4.0, 7.0],
        caption="Таблица 3. Декомпозиция при нормализации")


def section_er(doc: Document) -> None:
    add_h1(doc, "4. ER-диаграмма")
    add_body(doc,
        "Логическая модель данных представлена в нотации IDEF1x. Прямоугольник "
        "обозначает сущность, в верхней части — первичный ключ (PK), ниже — "
        "неключевые атрибуты. Связи «один-ко-многим» показаны линией; со стороны "
        "«многие» ставится точка. Идентифицирующие связи (дочерняя сущность не "
        "существует без родителя) показаны сплошной линией.")
    add_body(doc,
        "Перечень связей между сущностями приведён в таблице 4. Графическое "
        "изображение модели представлено на рисунке 1.",
        first_line_indent=False)
    if ER_IMAGE.exists():
        doc.add_picture(str(ER_IMAGE), width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph("Рисунок 1. ER-диаграмма базы данных")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(12)
    add_table_grid(doc,
        ["Родитель", "Потомок", "Тип", "Смысл связи"],
        [
            ["users", "dialogues", "1 : N", "у пользователя много диалогов"],
            ["dialogues", "messages", "1 : N", "в диалоге много сообщений"],
            ["messages", "attachments", "1 : N", "к сообщению несколько вложений"],
            ["dialogues", "dialogue_summaries", "1 : 1", "одна свёртка на диалог"],
            ["llm_providers", "llm_models", "1 : N", "у провайдера много моделей"],
            ["llm_models", "messages", "1 : N", "модель сгенерировала много ответов"],
            ["event_kinds", "analytics_events", "1 : N", "у типа события много записей"],
            ["users", "analytics_events", "1 : N", "события пользователя"],
            ["dialogues", "analytics_events", "1 : N", "события в контексте диалога"],
        ],
        widths=[3.4, 3.6, 2.0, 7.0],
        caption="Таблица 4. Связи между сущностями")
    add_body(doc,
        "Кардинальности и обязательность связей: связь users→dialogues "
        "идентифицирующая и обязательная со стороны диалога (диалог не может "
        "существовать без пользователя). Связи analytics_events→users и "
        "analytics_events→dialogues необязательные (ON DELETE SET NULL): "
        "системное событие может быть не привязано к пользователю, а удаление "
        "пользователя не должно стирать историю аналитики.")


def section_physical(doc: Document) -> None:
    add_h1(doc, "5. Физическая модель (DDL)")
    add_body(doc,
        "Ниже описаны ключевые таблицы и применённые ограничения целостности. "
        "Полный текст DDL приведён в Приложении А.")
    add_h2(doc, "5.1. Ограничения целостности")
    add_bullets(doc, [
        "PRIMARY KEY — суррогатные ключи BIGSERIAL/SMALLSERIAL для стабильности FK.",
        "FOREIGN KEY с явными политиками: CASCADE (диалоги/сообщения за "
        "владельцем), RESTRICT (нельзя удалить используемую модель), "
        "SET NULL (события переживают удаление пользователя).",
        "UNIQUE — telegram_id пользователя, коды провайдеров, моделей и событий.",
        "CHECK — допустимые роли сообщений, известные семейства моделей, "
        "неотрицательные токены и размеры, формат SHA-256, согласованность "
        "«assistant ⇒ есть модель».",
        "NOT NULL и DEFAULT — обязательные поля и значения по умолчанию "
        "(язык 'ru', статусы, временные метки NOW()).",
    ])
    add_h2(doc, "5.2. Пример определения таблицы messages")
    add_code(doc, _extract_messages_ddl(),
             caption="Листинг 2. Определение таблицы messages")
    add_body(doc,
        "Особый интерес представляет ограничение chk_messages_assistant_has_model: "
        "оно реализует бизнес-правило на уровне БД — у ответа модели обязательно "
        "указана модель, а у пользовательского сообщения её быть не должно. Это "
        "пример того, как инвариант предметной области защищается СУБД, а не "
        "только приложением.")


def _extract_messages_ddl() -> str:
    sql = read_sql("01_schema.sql")
    start = sql.find("CREATE TABLE messages")
    end = sql.find(");", start) + 2
    return sql[start:end]


def section_indexes(doc: Document) -> None:
    add_h1(doc, "6. Индексы")
    add_body(doc,
        "Индексы созданы под конкретные типовые запросы информационной системы. "
        "Назначение каждого индекса приведено в таблице 5. Полный текст — в "
        "Приложении Б.")
    add_table_grid(doc,
        ["Индекс", "Таблица / тип", "Под какой запрос"],
        [
            ["ix_dialogues_user_updated_at_desc", "dialogues, B-tree",
             "список диалогов пользователя по последней активности"],
            ["ix_dialogues_active_user", "dialogues, partial",
             "только неархивные диалоги пользователя"],
            ["ix_messages_dialogue_created", "messages, B-tree",
             "постраничное чтение сообщений диалога"],
            ["ix_messages_model_created", "messages, partial",
             "аналитика использования модели за период"],
            ["ix_attachments_message", "attachments, B-tree",
             "вложения сообщения (JOIN)"],
            ["ix_attachments_sha256", "attachments, B-tree",
             "дедупликация файлов по хешу"],
            ["ix_events_user_occurred", "analytics_events, B-tree",
             "события пользователя по времени"],
            ["ix_events_kind_occurred", "analytics_events, B-tree",
             "события одного типа по времени"],
            ["gin_events_properties", "analytics_events, GIN",
             "поиск по полям внутри JSONB (оператор @>)"],
        ],
        widths=[5.5, 4.0, 6.5],
        caption="Таблица 5. Индексы и их назначение")
    add_body(doc,
        "Частичные индексы (ix_dialogues_active_user, ix_messages_model_created) "
        "занимают меньше места и быстрее, поскольку покрывают только релевантное "
        "подмножество строк. GIN-индекс с классом jsonb_path_ops оптимизирован "
        "под оператор включения @> и компактнее стандартного jsonb_ops.")
    add_body(doc,
        "Проверить использование индекса можно командой EXPLAIN (ANALYZE, "
        "BUFFERS) перед запросом: план должен показывать Index Scan или Bitmap "
        "Index Scan вместо Seq Scan для выборок по индексированным условиям.")


def section_triggers(doc: Document) -> None:
    add_h1(doc, "7. Триггеры, функции и процедуры")

    add_h2(doc, "7.1. Триггеры")
    add_table_grid(doc,
        ["Триггер", "Событие", "Назначение"],
        [
            ["trg_dialogues_touch_updated_at", "BEFORE UPDATE dialogues",
             "поддерживать актуальный updated_at"],
            ["trg_messages_touch_dialogue", "AFTER INSERT messages",
             "обновлять updated_at диалога при новом сообщении"],
            ["trg_messages_log_event", "AFTER INSERT messages",
             "автоматически писать событие message_sent"],
            ["trg_messages_touch_user_last_seen", "AFTER INSERT messages",
             "обновлять last_seen_at пользователя"],
            ["trg_analytics_events_immutable", "BEFORE UPDATE/DELETE",
             "запрет удаления и ручных правок журнала (append-only)"],
        ],
        widths=[5.5, 4.5, 6.0],
        caption="Таблица 6. Триггеры")
    add_body(doc,
        "Триггер trg_analytics_events_immutable реализует неизменяемость журнала: "
        "удаление записей запрещено всегда, а из изменений разрешено единственное — "
        "системное обнуление внешних ключей при удалении связанного пользователя "
        "или диалога (ON DELETE SET NULL). Любая иная правка отклоняется. Это "
        "сохраняет достоверность аналитики на уровне СУБД и при этом не мешает "
        "штатному удалению связанных сущностей.")
    add_code(doc, _extract_block("03_triggers.sql",
             "CREATE OR REPLACE FUNCTION trg_fn_analytics_events_immutable",
             "EXECUTE FUNCTION trg_fn_analytics_events_immutable();"),
             caption="Листинг 3. Триггер неизменяемости журнала")

    add_h2(doc, "7.2. Функции")
    add_table_grid(doc,
        ["Функция", "Возвращает", "Назначение"],
        [
            ["fn_user_message_count", "BIGINT", "число сообщений пользователя за период"],
            ["fn_dialogue_token_total", "BIGINT", "суммарные токены диалога"],
            ["fn_top_models", "TABLE", "топ моделей по использованию за период"],
            ["fn_user_active_days", "INTEGER", "число активных дней (для DAU/MAU)"],
            ["fn_dialogue_message_count", "INTEGER", "число сообщений в диалоге"],
        ],
        widths=[5.0, 3.0, 8.0],
        caption="Таблица 7. Пользовательские функции")
    add_code(doc, _extract_block("04_functions.sql",
             "CREATE OR REPLACE FUNCTION fn_top_models",
             "$$;", first=True),
             caption="Листинг 4. Функция fn_top_models")

    add_h2(doc, "7.3. Хранимые процедуры")
    add_table_grid(doc,
        ["Процедура", "Назначение"],
        [
            ["sp_register_user", "UPSERT пользователя по telegram_id"],
            ["sp_create_dialogue", "создание нового диалога"],
            ["sp_compact_dialogue", "свёртка длинного диалога, удаление старых сообщений"],
            ["sp_archive_old_dialogues", "массовая архивация неактивных диалогов"],
            ["sp_log_event", "единая точка записи события в журнал"],
        ],
        widths=[5.0, 11.0],
        caption="Таблица 8. Хранимые процедуры")
    add_body(doc,
        "Процедура sp_compact_dialogue моделирует реальный механизм управления "
        "контекстом LLM: когда история превышает контекстное окно модели, старые "
        "сообщения заменяются их кратким пересказом (свёрткой), а сами строки "
        "удаляются. Процедура выполняет UPSERT свёртки, удаление устаревших "
        "сообщений и запись события в одной транзакции.")
    add_code(doc, _extract_block("05_procedures.sql",
             "CREATE OR REPLACE PROCEDURE sp_compact_dialogue",
             "$$;", first=True),
             caption="Листинг 5. Процедура sp_compact_dialogue")

    add_h2(doc, "7.4. О хранении свойств событий в JSONB")
    add_body(doc,
        "Журнал analytics_events содержит поле properties_json типа JSONB. "
        "Набор свойств различается у разных типов событий и часто меняется при "
        "развитии продукта. Возможные альтернативы и их недостатки:")
    add_bullets(doc, [
        "Отдельная таблица под каждый тип события — десятки таблиц и миграция "
        "при каждом новом событии.",
        "Одна «широкая» таблица со всеми возможными полями — сильная "
        "разрежённость (большинство полей NULL).",
        "Модель «сущность-атрибут-значение» (EAV) — громоздкие запросы и "
        "потеря типобезопасности.",
    ])
    add_body(doc,
        "JSONB с GIN-индексом — компромисс: схема гибкая, поиск по ключам "
        "(@>) индексируется, при этом ядро модели (кто, когда, какой тип "
        "события) остаётся строго реляционным через FK на users, dialogues и "
        "event_kinds.")


def _extract_block(filename: str, start_marker: str, end_marker: str,
                  first: bool = False) -> str:
    sql = read_sql(filename)
    start = sql.find(start_marker)
    if start == -1:
        return "-- (фрагмент не найден)"
    if first:
        end = sql.find(end_marker, start) + len(end_marker)
    else:
        end = sql.find(end_marker, start) + len(end_marker)
    return sql[start:end].strip()


def section_nosql(doc: Document) -> None:
    add_h1(doc, "8. Оценка целесообразности NoSQL")
    add_body(doc,
        "Под NoSQL понимают нереляционные хранилища: документные (MongoDB), "
        "ключ-значение (Redis), колоночные (Cassandra), графовые (Neo4j). Оценим, "
        "даст ли переход выгоду для данной предметной области.")
    add_h2(doc, "8.1. Где NoSQL был бы уместен")
    add_bullets(doc, [
        "Документное хранилище (MongoDB) органично хранило бы диалог как один "
        "документ со встроенным массивом сообщений — удобно для чтения всей "
        "переписки одним запросом.",
        "Кэш ключ-значение (Redis) хорошо подошёл бы для горячих данных: "
        "активная сессия, статус «модель печатает», счётчики.",
        "При огромных объёмах журнала событий колоночное хранилище "
        "масштабируется по записи лучше реляционного.",
    ])
    add_h2(doc, "8.2. Где NoSQL проигрывает")
    add_bullets(doc, [
        "Нет внешних ключей и каскадов — целостность связей пришлось бы "
        "поддерживать в приложении, что повышает риск рассинхронизации.",
        "Аналитика по связанным сущностям (пользователи × модели × периоды) "
        "требует JOIN и агрегаций, которые в документных БД выражаются "
        "громоздко и работают медленнее.",
        "Дублирование данных пользователя/модели внутри каждого документа "
        "возвращает аномалии обновления, от которых избавила нормализация.",
        "Транзакционная согласованность (например, свёртка диалога) "
        "реализуется в реляционной БД проще и надёжнее.",
    ])
    add_h2(doc, "8.3. Вывод")
    add_body(doc,
        "Для данной задачи ядром системы должна оставаться реляционная СУБД "
        "PostgreSQL: предметная область сильно связана и аналитична. NoSQL "
        "оправдан лишь как дополнение — Redis для кэша и, при кратном росте "
        "объёма событий, отдельное колоночное хранилище для журнала. Полный "
        "переход на NoSQL ухудшил бы целостность и усложнил аналитику, поэтому "
        "нецелесообразен.")


def section_crud(doc: Document) -> None:
    add_h1(doc, "9. Реализация CRUD через REST API")
    add_body(doc,
        "Для демонстрации взаимодействия с базой данных реализовано веб-приложение "
        "на Python (FastAPI) с использованием ORM SQLAlchemy. Приложение "
        "предоставляет REST API с операциями CRUD над основными сущностями. "
        "Исходный код находится в каталоге app/, инструкция по запуску — в "
        "app/README.md.")
    add_table_grid(doc,
        ["Метод и путь", "Операция CRUD", "Действие"],
        [
            ["POST /users", "Create", "регистрация пользователя"],
            ["GET /users/{id}", "Read", "получить пользователя"],
            ["GET /users", "Read", "список пользователей"],
            ["POST /dialogues", "Create", "создать диалог"],
            ["GET /dialogues?user_id=", "Read", "диалоги пользователя"],
            ["PATCH /dialogues/{id}", "Update", "переименовать/архивировать"],
            ["DELETE /dialogues/{id}", "Delete", "удалить диалог"],
            ["POST /dialogues/{id}/messages", "Create", "добавить сообщение"],
            ["GET /dialogues/{id}/messages", "Read", "сообщения диалога"],
            ["GET /analytics/top-models", "Read", "аналитика: топ моделей"],
        ],
        widths=[6.0, 3.0, 7.0],
        caption="Таблица 9. Эндпоинты REST API")
    add_body(doc,
        "Каждый эндпоинт транслируется в SQL-операции через ORM; аналитические "
        "эндпоинты вызывают созданные ранее функции БД (например, fn_top_models), "
        "что переиспользует серверную логику.")


def section_conclusion(doc: Document) -> None:
    add_h1(doc, "Заключение")
    add_body(doc,
        "В ходе проекта спроектирована и реализована реляционная база данных для "
        "предметной области «История диалогов и аналитика AI-чат-сервиса». "
        "Схема из девяти таблиц приведена к третьей нормальной форме, что "
        "исключает аномалии вставки, обновления и удаления.")
    add_body(doc,
        "Реализованы ограничения целостности (PK, FK с политиками CASCADE/"
        "RESTRICT/SET NULL, UNIQUE, CHECK), девять индексов под типовые запросы, "
        "пять триггеров (включая гарантию неизменяемости журнала), пять функций "
        "и пять хранимых процедур. Выполнена оценка целесообразности NoSQL: "
        "для связанной аналитической предметной области реляционная СУБД "
        "PostgreSQL предпочтительнее, NoSQL уместен лишь как дополнение.")
    add_body(doc,
        "Дополнительно реализовано CRUD-взаимодействие с БД через REST API на "
        "FastAPI и SQLAlchemy. Поставленные задачи выполнены, цель проекта "
        "достигнута.")


def section_sources(doc: Document) -> None:
    add_h1(doc, "Список источников")
    items = [
        "PostgreSQL 16 Documentation. — URL: https://www.postgresql.org/docs/16/ "
        "(дата обращения: 08.06.2026).",
        "Дейт К. Дж. Введение в системы баз данных. — М.: Вильямс, 2019.",
        "Кузнецов С. Д. Основы баз данных. — М.: БИНОМ, 2012.",
        "PostgreSQL: JSON Types and Functions. — URL: "
        "https://www.postgresql.org/docs/16/datatype-json.html "
        "(дата обращения: 08.06.2026).",
        "FastAPI Documentation. — URL: https://fastapi.tiangolo.com/ "
        "(дата обращения: 08.06.2026).",
        "SQLAlchemy 2.0 Documentation. — URL: https://docs.sqlalchemy.org/ "
        "(дата обращения: 08.06.2026).",
    ]
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(f"{i}. {it}")


def appendix_sql(doc: Document) -> None:
    doc.add_page_break()
    add_h1(doc, "Приложение А. Полный текст DDL (схема)")
    add_code(doc, read_sql("01_schema.sql"))

    doc.add_page_break()
    add_h1(doc, "Приложение Б. Индексы, триггеры, функции, процедуры")
    add_h2(doc, "Б.1. Индексы")
    add_code(doc, read_sql("02_indexes.sql"))
    add_h2(doc, "Б.2. Триггеры")
    add_code(doc, read_sql("03_triggers.sql"))
    add_h2(doc, "Б.3. Функции")
    add_code(doc, read_sql("04_functions.sql"))
    add_h2(doc, "Б.4. Процедуры")
    add_code(doc, read_sql("05_procedures.sql"))

    doc.add_page_break()
    add_h1(doc, "Приложение В. Код ER-диаграммы (для отрисовки)")
    add_body(doc,
        "Ниже — описание схемы в синтаксисе dbdiagram.io (DBML). Вставьте на "
        "dbdiagram.io, чтобы получить графическую ER-диаграмму для печати.",
        first_line_indent=False)
    add_code(doc, _dbml(), caption="Листинг 6. DBML-описание схемы")

    doc.add_page_break()
    add_h1(doc, "Приложение Г. Типовые SQL-запросы")
    add_code(doc, read_sql("07_sample_queries.sql"))


def _dbml() -> str:
    return (
        "Table users {\n"
        "  id bigint [pk]\n"
        "  telegram_id bigint [unique, not null]\n"
        "  username varchar\n"
        "  first_name varchar\n"
        "  last_name varchar\n"
        "  language_code varchar\n"
        "  registered_at timestamptz\n"
        "  last_seen_at timestamptz\n"
        "}\n\n"
        "Table llm_providers {\n"
        "  id smallint [pk]\n"
        "  code varchar [unique]\n"
        "  display_name varchar\n"
        "}\n\n"
        "Table llm_models {\n"
        "  id smallint [pk]\n"
        "  provider_id smallint [ref: > llm_providers.id]\n"
        "  code varchar [unique]\n"
        "  display_name varchar\n"
        "  family varchar\n"
        "  is_premium boolean\n"
        "  is_active boolean\n"
        "}\n\n"
        "Table dialogues {\n"
        "  id bigint [pk]\n"
        "  user_id bigint [ref: > users.id]\n"
        "  title varchar\n"
        "  created_at timestamptz\n"
        "  updated_at timestamptz\n"
        "  is_archived boolean\n"
        "}\n\n"
        "Table messages {\n"
        "  id bigint [pk]\n"
        "  dialogue_id bigint [ref: > dialogues.id]\n"
        "  role varchar\n"
        "  content text\n"
        "  model_id smallint [ref: > llm_models.id]\n"
        "  tokens_in int\n"
        "  tokens_out int\n"
        "  created_at timestamptz\n"
        "}\n\n"
        "Table attachments {\n"
        "  id bigint [pk]\n"
        "  message_id bigint [ref: > messages.id]\n"
        "  kind varchar\n"
        "  file_name varchar\n"
        "  mime_type varchar\n"
        "  size_bytes bigint\n"
        "  sha256 char\n"
        "}\n\n"
        "Table dialogue_summaries {\n"
        "  dialogue_id bigint [pk, ref: - dialogues.id]\n"
        "  summary text\n"
        "  summarized_until_message_id bigint [ref: > messages.id]\n"
        "  updated_at timestamptz\n"
        "}\n\n"
        "Table event_kinds {\n"
        "  id smallint [pk]\n"
        "  code varchar [unique]\n"
        "  description varchar\n"
        "}\n\n"
        "Table analytics_events {\n"
        "  id bigint [pk]\n"
        "  user_id bigint [ref: > users.id]\n"
        "  dialogue_id bigint [ref: > dialogues.id]\n"
        "  event_kind_id smallint [ref: > event_kinds.id]\n"
        "  source varchar\n"
        "  occurred_at timestamptz\n"
        "  properties_json jsonb\n"
        "}\n"
    )


def signatures(doc: Document) -> None:
    doc.add_page_break()
    add_h1(doc, "Лист подписей разработчиков")
    add_body(doc,
        "Каждый разработчик подтверждает участие в проекте и готовность "
        "отвечать на вопросы по всем его разделам.", first_line_indent=False)
    add_table_grid(doc,
        ["ФИО", "Раздел ответственности", "Подпись"],
        [
            ["Шарманов Даниил Андреевич", "Схема, нормализация, индексы", ""],
            ["", "Триггеры, функции, процедуры", ""],
            ["", "ER-диаграмма, оценка NoSQL", ""],
            ["", "CRUD REST API, тестовые данные", ""],
        ],
        widths=[6.0, 7.0, 3.0])


def main() -> None:
    doc = Document()
    style_base(doc)
    # поля страницы
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    title_page(doc)
    section_intro(doc)
    section_subject(doc)
    section_dbms(doc)
    section_norm(doc)
    section_er(doc)
    section_physical(doc)
    section_indexes(doc)
    section_triggers(doc)
    section_nosql(doc)
    section_crud(doc)
    section_conclusion(doc)
    section_sources(doc)
    appendix_sql(doc)
    signatures(doc)

    doc.core_properties.title = "Проектирование базы данных истории диалогов и аналитики AI-чат-сервиса"
    doc.core_properties.author = "Шарманов Даниил Андреевич"
    doc.core_properties.subject = "Курсовой проект по дисциплине «Базы данных»"
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
